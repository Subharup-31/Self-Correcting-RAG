"""
Document loader for mixed, messy document sets.

Handles:
  - Digital PDFs         (pymupdf text extraction)
  - Scanned image PDFs   (pytesseract OCR fallback)
  - Standalone images    (pytesseract OCR)
  - HTML                 (unstructured)
  - Word .docx           (python-docx via unstructured)
  - Plain text / Markdown

Each extracted page becomes a LangChain `Document` tagged with metadata:
  source, page_number, doc_type, ocr_confidence, char_count, source_file

OCR heuristic: if pymupdf extracts < 50 chars from a page, treat the page as
scanned and run pytesseract. Average per-word confidence is recorded so the
pipeline can down-weight low-quality OCR chunks later.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

from langchain_core.documents import Document
from loguru import logger


# Character threshold below which a PDF page is considered "scanned".
DIGITAL_TEXT_THRESHOLD = 50


class DocumentLoader:
    """Loads messy documents into clean, metadata-tagged LangChain Documents."""

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif",
        ".html", ".htm", ".docx", ".doc", ".txt", ".md",
    }

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #
    def load(self, path: str) -> List[Document]:
        """Load a single file into one or more Documents (one per page)."""
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"Document not found: {p}")

        ext = p.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"Unsupported file type {ext} for {p.name}; skipping.")
            return []

        logger.info(f"Loading document: {p.name} ({ext})")
        try:
            if ext == ".pdf":
                docs = self._load_pdf(p)
            elif ext in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"}:
                docs = self._load_image(p)
            elif ext in {".html", ".htm"}:
                docs = self._load_html(p)
            elif ext in {".docx", ".doc"}:
                docs = self._load_word(p)
            else:  # .txt, .md
                docs = self._load_text(p)
        except Exception as exc:  # noqa: BLE001
            logger.exception(f"Failed to load {p.name}: {exc}")
            return []

        logger.info(f"Loaded {len(docs)} page(s) from {p.name}")
        return docs

    def load_directory(self, dir_path: str) -> List[Document]:
        """Recursively load every supported file in a directory."""
        d = Path(dir_path)
        if not d.exists() or not d.is_dir():
            logger.warning(f"Directory not found: {d}")
            return []

        all_docs: List[Document] = []
        for p in sorted(d.rglob("*")):
            if p.is_file() and not p.name.startswith("."):
                if p.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                    all_docs.extend(self.load(str(p)))
        logger.info(f"Loaded {len(all_docs)} total pages from {d}")
        return all_docs

    # ------------------------------------------------------------------ #
    # PDF — the most complex case (digital + scanned handling)
    # ------------------------------------------------------------------ #
    def _load_pdf(self, path: Path) -> List[Document]:
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:  # pragma: no cover
            raise ImportError(
                "PyMuPDF (fitz) is required for PDF loading. "
                "Install with: pip install pymupdf"
            ) from exc

        docs: List[Document] = []
        with fitz.open(str(path)) as doc:
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text") or ""
                text = self._clean_text(text)

                if len(text.strip()) >= DIGITAL_TEXT_THRESHOLD:
                    # Digital page — text extraction succeeded.
                    docs.append(Document(
                        page_content=text,
                        metadata=self._meta(
                            path, page_num, "pdf_digital", ocr_confidence=1.0,
                            char_count=len(text),
                        ),
                    ))
                else:
                    # Scanned / image-only page → OCR fallback.
                    ocr_text, conf = self._ocr_page(page)
                    ocr_text = self._clean_text(ocr_text)
                    docs.append(Document(
                        page_content=ocr_text,
                        metadata=self._meta(
                            path, page_num, "pdf_scanned",
                            ocr_confidence=conf, char_count=len(ocr_text),
                        ),
                    ))
        return docs

    def _ocr_page(self, page) -> Tuple[str, float]:
        """Render a fitz page to an image and run pytesseract OCR.

        Returns (text, mean_word_confidence). Confidence is 0.0 if OCR is
        unavailable or fails.
        """
        try:
            import pytesseract
            from PIL import Image
            import io
        except ImportError:  # pragma: no cover
            logger.warning("pytesseract/PIL not installed; OCR unavailable.")
            return "", 0.0

        try:
            # Render page at 300 DPI for decent OCR accuracy.
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            data = pytesseract.image_to_data(
                img, output_type=pytesseract.Output.DICT
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning(f"OCR failed on page {getattr(page, 'number', '?')}: {exc}")
            return "", 0.0

        # Reconstruct text and compute mean confidence over recognized words.
        words, confs = [], []
        for i, w in enumerate(data.get("text", [])):
            w = (w or "").strip()
            if not w:
                continue
            try:
                conf = float(data["conf"][i])
            except (ValueError, KeyError, IndexError):
                conf = -1.0
            if conf >= 0:  # -1 means non-text block
                words.append(w)
                confs.append(conf)
            else:
                words.append(w)

        text = " ".join(words)
        mean_conf = sum(confs) / len(confs) / 100.0 if confs else 0.0
        return text, round(mean_conf, 3)

    # ------------------------------------------------------------------ #
    # Image (standalone)
    # ------------------------------------------------------------------ #
    def _load_image(self, path: Path) -> List[Document]:
        try:
            import pytesseract
            from PIL import Image
        except ImportError as exc:  # pragma: no cover
            raise ImportError(
                "pytesseract + Pillow required for image OCR."
            ) from exc

        img = Image.open(path)
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        words, confs = [], []
        for i, w in enumerate(data.get("text", [])):
            w = (w or "").strip()
            if not w:
                continue
            try:
                conf = float(data["conf"][i])
            except (ValueError, KeyError, IndexError):
                conf = -1.0
            words.append(w)
            if conf >= 0:
                confs.append(conf)
        text = self._clean_text(" ".join(words))
        mean_conf = sum(confs) / len(confs) / 100.0 if confs else 0.0

        return [Document(
            page_content=text,
            metadata=self._meta(path, 1, "image", ocr_confidence=mean_conf,
                                char_count=len(text)),
        )]

    # ------------------------------------------------------------------ #
    # HTML
    # ------------------------------------------------------------------ #
    def _load_html(self, path: Path) -> List[Document]:
        try:
            from unstructured.partition.html import partition_html
        except ImportError:  # pragma: no cover
            # Fallback: strip tags with regex.
            raw = path.read_text(encoding="utf-8", errors="ignore")
            text = self._clean_text(re.sub(r"<[^>]+>", " ", raw))
            return [Document(page_content=text,
                             metadata=self._meta(path, 1, "html", 1.0, len(text)))]

        elements = partition_html(filename=str(path))
        text = self._clean_text("\n".join(str(e) for e in elements))
        return [Document(page_content=text,
                         metadata=self._meta(path, 1, "html", 1.0, len(text)))]

    # ------------------------------------------------------------------ #
    # Word
    # ------------------------------------------------------------------ #
    def _load_word(self, path: Path) -> List[Document]:
        try:
            from unstructured.partition.docx import partition_docx
        except ImportError:  # pragma: no cover
            try:
                from docx import Document as DocxDocument
                d = DocxDocument(str(path))
                text = self._clean_text("\n".join(p.text for p in d.paragraphs))
                return [Document(page_content=text,
                                 metadata=self._meta(path, 1, "word", 1.0, len(text)))]
            except ImportError as exc:
                raise ImportError(
                    "python-docx or unstructured required for .docx loading."
                ) from exc

        elements = partition_docx(filename=str(path))
        text = self._clean_text("\n".join(str(e) for e in elements))
        return [Document(page_content=text,
                         metadata=self._meta(path, 1, "word", 1.0, len(text)))]

    # ------------------------------------------------------------------ #
    # Plain text / Markdown
    # ------------------------------------------------------------------ #
    def _load_text(self, path: Path) -> List[Document]:
        # Try common encodings.
        text = None
        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                text = path.read_text(encoding=enc)
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            text = path.read_bytes().decode("utf-8", errors="ignore")
        text = self._clean_text(text)
        return [Document(page_content=text,
                         metadata=self._meta(path, 1, "text", 1.0, len(text)))]

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #
    def _clean_text(self, text: str) -> str:
        """Normalize whitespace, drop common header/footer artifacts, fix
        encoding gremlins."""
        if not text:
            return ""
        # Replace common mojibake / smart quotes.
        replacements = {
            "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
            "\u2013": "-", "\u2014": "-", "\u00a0": " ", "\ufb01": "fi",
            "\ufb02": "fl", "\u2026": "...",
        }
        for bad, good in replacements.items():
            text = text.replace(bad, good)

        # Collapse runs of whitespace, preserving paragraph breaks.
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Strip obvious page-header/footer artifacts (e.g. "Page 3 of 12").
        text = re.sub(r"(?im)^\s*page\s+\d+\s+(of|/)\s+\d+\s*$", "", text)

        return text.strip()

    @staticmethod
    def _meta(path: Path, page_number: int, doc_type: str,
              ocr_confidence: float, char_count: int) -> dict:
        return {
            "source": str(path.name),
            "source_path": str(path),
            "page_number": page_number,
            "doc_type": doc_type,
            "ocr_confidence": float(ocr_confidence),
            "char_count": int(char_count),
        }


def load_document(path: str) -> List[Document]:
    """Convenience function: load a single document."""
    return DocumentLoader().load(path)


def load_directory(path: str) -> List[Document]:
    """Convenience function: load all documents in a directory."""
    return DocumentLoader().load_directory(path)
